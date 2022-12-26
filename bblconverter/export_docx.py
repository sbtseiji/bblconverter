import re
from docx import Document
from bblconverter.export_cleaner import cleaner

out_str = []
def export_docx(biblist, outfile):
  doc = Document()

  i = 0
  for bib in biblist:
    res_str = ''.join(bib[1])
    res = cleaner(res_str,'docx')
    match_obj = re.match(r'(.*?)(<italic>|<bold>)(.*?)(</italic>|</bold>)', res)

    if match_obj is None:
      doc.add_paragraph(res)

    else:
      doc.add_paragraph()

      while not match_obj is None:
        doc.paragraphs[i].add_run(match_obj[1])
        if match_obj[2] == '<italic>':
          doc.paragraphs[i].add_run(match_obj[3]).italic = True
        else:
          doc.paragraphs[i].add_run(match_obj[3]).bold = True

        res = res.replace(match_obj[0], '', 1)
        match_obj = re.match(r'(.*?)(<italic>|<bold>)(.*?)(</italic>|</bold>)', res)

        if match_obj is None:
          doc.paragraphs[i].add_run(res)
    i += 1

  doc.save(outfile)